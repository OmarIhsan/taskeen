import openpyxl
import os
from docx import Document
from docx.oxml import parse_xml
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import tkinter as tk
from tkinter import filedialog
    

def distribute_names(excel_file, sheet_name, output_folder):
    try:
        # Load Excel workbook and worksheet
        workbook = openpyxl.load_workbook(excel_file)
        worksheet = workbook[sheet_name]

        # Define column references
        name_col = 'A'
        group_size_col = 'H'
        group_name_col = 'M'
        gender_col = 'N'
        k_col = 'K'

        # Initialize variables
        used_names = set()
        document = Document()
        group_num = 1
        tables_buffer = []

        # Set page orientation to landscape
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        
        # Set two columns
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num', '2')

        # Set page margins to be smaller
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        # Set default font size
        style = document.styles['Normal']
        style.font.size = Pt(18)

        def create_merged_table():
            if tables_buffer:
                table = document.add_table(rows=6, cols=3)
                table.style = 'Table Grid'
                table.autofit = True
                
                for col_idx, data in enumerate(tables_buffer):
                    for row_idx in range(6):
                        cell = table.cell(row_idx, col_idx)
                        if row_idx < len(data):
                            cell.text = data[row_idx]
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if row_idx == 0:
                                cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="FFFF00"/>'))
                                cell.paragraphs[0].runs[0].bold = True
                                cell.paragraphs[0].runs[0].font.size = Pt(20)
                
                document.add_paragraph().paragraph_format.space_after = Pt(18)
                tables_buffer.clear()

        # Process each row in the worksheet
        for row in range(2, worksheet.max_row + 1):
            group_name = worksheet[f'{group_name_col}{row}'].value
            name = worksheet[f'{name_col}{row}'].value
            group_size = worksheet[f'{group_size_col}{row}'].value

            

            if name and name not in used_names:
                current_table_data = []
                
                if group_size == 1:
                    current_table_data.append(f"{group_name} فردي" if group_name else "فردي شركة")
                    k_value = worksheet[f'{k_col}{row}'].value
                    name_text = f"{name} $" if k_value else name
                    current_table_data.append(name_text)
                    used_names.add(name)
                
                elif group_size == 2:
                    next_row = row + 1
                    if next_row <= worksheet.max_row:
                        next_name = worksheet[f'{name_col}{next_row}'].value
                        if next_name and next_name not in used_names:
                            current_table_data.append(f"{group_name} ثنائي" if group_name else "ثنائي شركة")
                            k_value = worksheet[f'{k_col}{row}'].value
                            name_text = f"{name} $" if k_value else name
                            current_table_data.append(name_text)
                            k_value_next = worksheet[f'{k_col}{next_row}'].value
                            next_name_text = f"{next_name} $" if k_value_next else next_name
                            current_table_data.append(next_name_text)
                            used_names.add(name)
                            used_names.add(next_name)
                
                elif group_size == 3:
                    next_row = row + 1
                    next_row2 = row + 2
                    if next_row2 <= worksheet.max_row:
                        next_name = worksheet[f'{name_col}{next_row}'].value
                        next_name2 = worksheet[f'{name_col}{next_row2}'].value
                        if next_name and next_name2 and next_name not in used_names and next_name2 not in used_names:
                            current_table_data.append(f"{group_name} ثلاثي" if group_name else "ثلاثي شركة")
                            k_value = worksheet[f'{k_col}{row}'].value
                            name_text = f"{name} $" if k_value else name
                            current_table_data.append(name_text)
                            k_value_next = worksheet[f'{k_col}{next_row}'].value
                            next_name_text = f"{next_name} $" if k_value_next else next_name
                            current_table_data.append(next_name_text)
                            k_value_next2 = worksheet[f'{k_col}{next_row2}'].value
                            next_name2_text = f"{next_name2} $" if k_value_next2 else next_name2
                            current_table_data.append(next_name2_text)
                            used_names.add(name)
                            used_names.add(next_name)
                            used_names.add(next_name2)

                elif group_size == 4:
                    names_in_group = []
                    female_names = []
                    
                    for next_row in range(row, worksheet.max_row + 1):
                        next_name = worksheet[f'{name_col}{next_row}'].value
                        next_size = worksheet[f'{group_size_col}{next_row}'].value
                        gender = worksheet[f'{gender_col}{next_row}'].value
                        k_value = worksheet[f'{k_col}{next_row}'].value
                        group_name = worksheet[f'M{next_row}'].value
                        
                        if next_name and next_size == 4 and next_name not in used_names:
                            name_with_k = f"{next_name} $" if k_value else next_name
                            names_in_group.append((name_with_k, next_name, group_name))
                    
                    # Group by M column value
                    group_dict = {}
                    for name_with_k, name, group_name in names_in_group:
                        if group_name not in group_dict:
                            group_dict[group_name] = []
                        group_dict[group_name].append((name_with_k, name))
                    
                    # Process each group
                    for group_name, members in group_dict.items():
                        for i in range(0, len(members), 4):
                            current_group = members[i:i+4]
                            if len(current_group) == 4:  # Only process complete groups of 4
                                current_table_data = [f"رباعي {group_name}" if group_name else "رباعي شركة"]
                                for member_name_with_k, member_name in current_group:
                                    current_table_data.append(member_name_with_k)
                                    used_names.add(member_name)
                                
                                tables_buffer.append(current_table_data)
                                if len(tables_buffer) == 3:
                                    create_merged_table()

                            else:
                                remaining_members = members[i:]
                                for j in range(0, len(remaining_members), 2):
                                    current_group = remaining_members[j:j+2]
                                    current_table_data = [f"رباعي {group_name}" if group_name else "رباعي شركة"]
                                    for member_name_with_k, member_name in current_group:
                                        current_table_data.append(member_name_with_k)
                                        used_names.add(member_name)
                                    tables_buffer.append(current_table_data)
                                    if len(tables_buffer) == 3:
                                        create_merged_table()
                    
                elif group_size == 0:         
                    names_in_group = []
                    
                    for next_row in range(row, worksheet.max_row + 1):
                        next_name = worksheet[f'{name_col}{next_row}'].value
                        next_size = worksheet[f'{group_size_col}{next_row}'].value
                        k_value = worksheet[f'{k_col}{next_row}'].value
                        
                        if next_name and next_size == 0 and next_name not in used_names:
                            name_with_k = f"{next_name} $" if k_value else next_name
                            names_in_group.append((name_with_k, next_name))
                    
                    if names_in_group:
                        current_table_data = ["طفل"]
                        for member_name_with_k, member_name in names_in_group:
                            current_table_data.append(member_name_with_k)
                            used_names.add(member_name)
                        
                        tables_buffer.append(current_table_data)
                        if len(tables_buffer) == 3:
                            create_merged_table()                
                            
                elif group_size == 5:
                    names_in_group = []
                    female_names = []
                    
                    for next_row in range(row, worksheet.max_row + 1):
                        next_name = worksheet[f'{name_col}{next_row}'].value
                        next_size = worksheet[f'{group_size_col}{next_row}'].value
                        gender = worksheet[f'{gender_col}{next_row}'].value
                        k_value = worksheet[f'{k_col}{next_row}'].value
                        
                        if next_name and next_size == 5 and next_name not in used_names:
                            name_with_k = f"{next_name} $" if k_value else next_name
                            if gender == "انثى":
                                female_names.append((name_with_k, next_name))
                            else:
                                names_in_group.append((name_with_k, next_name))
                    
                    # Process male/other groups
                    if names_in_group:
                        for i in range(0, len(names_in_group), 5):
                            current_group = names_in_group[i:i+5]
                            current_table_data = ["خماسي رجال"]
                            for member_name_with_k, member_name in current_group:
                                for r in range(1, worksheet.max_row + 1):
                                    if worksheet[f'{name_col}{r}'].value == member_name:
                                        group_name = worksheet[f'M{r}'].value
                                        if group_name:
                                            current_table_data.append(f"{member_name_with_k} / {group_name}")
                                        else:
                                            current_table_data.append(member_name_with_k)
                                        used_names.add(member_name)
                                        break
                            
                            tables_buffer.append(current_table_data)
                            if len(tables_buffer) == 3:
                                create_merged_table()
                    
                    # Process female groups
                    if female_names:
                        for i in range(0, len(female_names), 5):
                            current_group = female_names[i:i+5]
                            current_table_data = ["خماسي نساء"]
                            for member_name_with_k, member_name in current_group:
                                for r in range(1, worksheet.max_row + 1):
                                    if worksheet[f'{name_col}{r}'].value == member_name:
                                        group_name = worksheet[f'M{r}'].value
                                        if group_name:
                                            current_table_data.append(f"{member_name_with_k} / {group_name}")
                                        else:
                                            current_table_data.append(member_name_with_k)
                                        used_names.add(member_name)
                                        break
                            
                            tables_buffer.append(current_table_data)
                            if len(tables_buffer) == 3:
                                create_merged_table()
                
                if current_table_data and group_size < 4:
                    tables_buffer.append(current_table_data)
                    if len(tables_buffer) == 3:
                        create_merged_table()
        
        # Create final merged table if there are remaining tables
        if tables_buffer:
            create_merged_table()

        document.save(f"{output_folder}/groups.docx")

    except FileNotFoundError:
        print(f"Error: Excel file '{excel_file}' not found")
    except KeyError:
        print(f"Error: Sheet '{sheet_name}' not found in the Excel file")
    except Exception as e:
        print(f"An error occurred: {str(e)}")

def main():
    
    def select_file():
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if file_path:
            file_path_label.config(text=f"Selected file: {file_path}")
            global excel_file_path
            excel_file_path = file_path
            
    def start_program():
        if 'excel_file_path' in globals():
            sheet_name = "رحلة"
            output_folder = os.path.dirname(excel_file_path)
            distribute_names(excel_file_path, sheet_name, output_folder)
            root.destroy()
        else:
            file_path_label.config(text="Please select an Excel file first!")
    
    root = tk.Tk()
    root.title("Excel File Processor")
    root.geometry("400x200")
    
    select_button = tk.Button(root, text="Select Excel File", command=select_file)
    select_button.pack(pady=20)
    
    file_path_label = tk.Label(root, text="No file selected")
    file_path_label.pack(pady=10)
    
    start_button = tk.Button(root, text="Start Program", command=start_program)
    start_button.pack(pady=20)
    
    root.mainloop()
if __name__ == "__main__":
    main()